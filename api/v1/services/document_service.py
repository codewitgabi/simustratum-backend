import io
import uuid
from typing import Any

from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile, status
from pypdf import PdfReader
from qdrant_client import models
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from api.v1.models.document import Document, DocumentStatus
from api.v1.services.cloudinary_client import upload_document as upload_to_cloudinary
from api.v1.services.embedding_service import embed_document_chunks, embed_query
from api.v1.services.qdrant_client import get_qdrant_client
from api.v1.utils.config import config
from api.v1.utils.logger import get_logger

logger = get_logger("document_service")

CHUNK_SIZE = 800
CHUNK_OVERLAP = 100

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}
MAX_DOCUMENT_SIZE_BYTES = 15 * 1024 * 1024


class UnsupportedDocumentError(ValueError):
    pass


def extract_text(filename: str, content: bytes) -> str:
    extension = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if extension == ".pdf":
        reader = PdfReader(io.BytesIO(content))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)

    if extension == ".docx":
        doc = DocxDocument(io.BytesIO(content))
        return "\n\n".join(p.text for p in doc.paragraphs)

    if extension == ".txt":
        return content.decode("utf-8", errors="ignore")

    raise UnsupportedDocumentError(f"Unsupported document type: {extension or filename}")


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    words = text.split()
    if not words:
        return []

    chunks = []
    step = chunk_size - overlap
    for start in range(0, len(words), step):
        chunk = " ".join(words[start : start + chunk_size])
        if chunk.strip():
            chunks.append(chunk)
    return chunks


async def process_document(document_id: uuid.UUID, text: str, db: AsyncSession) -> None:
    """
    Chunks and embeds a document's text into Qdrant, then updates the Document row
    with the outcome. Embedding failures mark the document FAILED rather than
    raising — the upload endpoint has already returned a response by the time this
    runs, so there's no request left to propagate an exception to.
    """
    result = await db.execute(select(Document).where(Document.id == document_id))
    document = result.scalar_one_or_none()
    if document is None:
        return

    chunks = chunk_text(text)
    if not chunks:
        document.status = DocumentStatus.FAILED
        document.error_message = "No extractable text found in document"
        await db.commit()
        return

    try:
        client = get_qdrant_client()
        vectors = await embed_document_chunks(chunks)
        points = [
            models.PointStruct(
                id=str(uuid.uuid4()),
                vector=vector,
                payload={"document_id": str(document_id), "chunk_index": i, "text": chunk},
            )
            for i, (chunk, vector) in enumerate(zip(chunks, vectors))
        ]
        await client.upsert(collection_name=config.QDRANT_COLLECTION, points=points)
    except Exception:
        logger.exception("Failed to embed document", extra={"document_id": str(document_id)})
        document.status = DocumentStatus.FAILED
        document.error_message = "Failed to embed document"
        await db.commit()
        return

    document.status = DocumentStatus.READY
    document.chunk_count = len(chunks)
    await db.commit()


async def retrieve_relevant_chunks(document_id: uuid.UUID, query_text: str, limit: int = 4) -> list[str]:
    """
    Returns the most relevant chunks of a document for the given query. Returns an
    empty list (rather than raising) if Qdrant is unreachable or the document has no
    embedded chunks — this is supplementary grounding context, not a hard dependency
    for the turn loop to function.
    """
    try:
        client = get_qdrant_client()
        query_vector = await embed_query(query_text)
        response = await client.query_points(
            collection_name=config.QDRANT_COLLECTION,
            query=query_vector,
            query_filter=models.Filter(
                must=[models.FieldCondition(key="document_id", match=models.MatchValue(value=str(document_id)))]
            ),
            limit=limit,
        )
        return [point.payload["text"] for point in response.points if point.payload]
    except Exception:
        logger.exception("Failed to retrieve document context", extra={"document_id": str(document_id)})
        return []


async def create_document(user_id: Any, file: UploadFile, db: AsyncSession) -> Document:
    filename = file.filename or "document"
    extension = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if extension not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported document type. Allowed types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}",
        )

    content = await file.read()
    if len(content) > MAX_DOCUMENT_SIZE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Document exceeds the {MAX_DOCUMENT_SIZE_BYTES // (1024 * 1024)}MB size limit",
        )
    if not content:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Document is empty")

    file_url = upload_to_cloudinary(content, filename)

    document = Document(
        user_id=user_id,
        filename=filename,
        file_url=file_url,
        status=DocumentStatus.PROCESSING,
    )
    db.add(document)
    await db.commit()
    await db.refresh(document)

    try:
        text = extract_text(filename, content)
    except UnsupportedDocumentError as exc:
        document.status = DocumentStatus.FAILED
        document.error_message = str(exc)
        await db.commit()
        return document

    await process_document(document.id, text, db)
    await db.refresh(document)
    return document


async def get_owned_document(user_id: Any, document_id: uuid.UUID, db: AsyncSession) -> Document | None:
    result = await db.execute(
        select(Document).where(Document.id == document_id, Document.user_id == user_id)
    )
    return result.scalar_one_or_none()
